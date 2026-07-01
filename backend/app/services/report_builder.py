"""
Bower Ag CowCare Tool — DOCX Report Builder
Sprint 9/21: Builds customer-facing reports using python-docx.

Uses the official Bower Ag letterhead logo centered at top.
Document B Section 5 structure: Cover, intro, findings, recommendations,
pricing table, next steps, about Bower Ag.

⚠️ Pricing table only included when include_pricing=True and pricing_table not empty.
"""

import io
import os
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ─── Constants ────────────────────────────────────────────────────────────────

NAVY = RGBColor(0x0D, 0x1F, 0x3C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_LIGHT = RGBColor(0xF3, 0xF4, 0xF6)
ACCENT_GREEN = RGBColor(0x4C, 0xAF, 0x50)
FONT_BODY = "Calibri"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_HEADER = Pt(14)

# Logo path — relative to this file
ASSETS_DIR = Path(__file__).parent.parent / "assets"
LOGO_PATH = ASSETS_DIR / "bower_ag_logo.jpg"

# Footer text from the sample report
FOOTER_TEXT = "The Dairy Solutions Group PO Box 3640 Turlock, California 95381 P 209-669-6200 TDSG.US"


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str) -> None:
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_section_header(doc: Document, text: str) -> None:
    """Add a navy section header (14pt bold Calibri)."""
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = FONT_SIZE_HEADER
    run.font.color.rgb = NAVY
    run.font.name = FONT_BODY


def _add_body_paragraph(doc: Document, text: str) -> None:
    """Add a standard body paragraph (11pt Calibri)."""
    p = doc.add_paragraph()
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = FONT_SIZE_BODY
    run.font.name = FONT_BODY


def _add_horizontal_rule(doc: Document) -> None:
    """Add a thin horizontal rule (bottom border on empty paragraph)."""
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def _parse_report_sections(report_content: str) -> dict[str, str]:
    """
    Parse Claude-generated report text into named sections.
    Looks for ## headers or **Header** patterns from Document B Section 5.
    """
    sections: dict[str, str] = {}

    # Normalize header patterns
    header_map = {
        "a quick note before we start": "intro",
        "what we found": "findings",
        "our recommendations": "recommendations",
        "your program summary": "program_summary",
        "what happens next": "next_steps",
        "about bower ag": "about",
    }

    # Split on ## headers or **bold** headers
    parts = re.split(r'(?:^|\n)(?:##\s*|\*\*)(.*?)(?:\*\*)?(?:\n)', report_content, flags=re.IGNORECASE)

    current_key: Optional[str] = None
    for i, part in enumerate(parts):
        part_stripped = part.strip()
        part_lower = part_stripped.lower().rstrip(":").strip("*# ")

        # Check if this part is a header
        matched_key = None
        for pattern, key in header_map.items():
            if pattern in part_lower:
                matched_key = key
                break

        if matched_key:
            current_key = matched_key
            sections[current_key] = ""
        elif current_key and part_stripped:
            sections[current_key] = (sections.get(current_key, "") + "\n" + part_stripped).strip()

    # If parsing fails, put everything in "body"
    if not sections:
        sections["body"] = report_content.strip()

    return sections


# ─── Main Builder ─────────────────────────────────────────────────────────────

def build_report_docx(
    customer_name: str,
    operation_name: str,
    location_name: str,
    rep_name: str,
    rep_title: str,
    report_date: str,
    report_content: str,
    pricing_table: list[dict],
    include_pricing: bool,
) -> bytes:
    """
    Build a DOCX report matching Document B Section 5 structure.

    Returns the DOCX file as bytes.
    """
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── Header (pages after cover) ──
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run(f"Bower Ag — {operation_name}")
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    hr.font.name = FONT_BODY

    # ── Footer ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(FOOTER_TEXT)
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    fr.font.name = FONT_BODY

    # ═══════════════════════════════════════════════════════════════════════════
    # PAGE 1: COVER WITH LOGO LETTERHEAD
    # ═══════════════════════════════════════════════════════════════════════════

    # Centered Bower Ag logo
    if LOGO_PATH.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(2.0))
    else:
        # Fallback: text logo if image not found
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_logo = p_logo.add_run("BOWER AG")
        r_logo.bold = True
        r_logo.font.size = Pt(36)
        r_logo.font.color.rgb = NAVY
        r_logo.font.name = FONT_BODY

    # Customer operation name — main title
    p_op = doc.add_paragraph()
    p_op.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_op.space_before = Pt(12)
    r_op = p_op.add_run(f"{customer_name} – Cow Care Program Summary")
    r_op.bold = True
    r_op.font.size = Pt(18)
    r_op.font.color.rgb = NAVY
    r_op.font.name = FONT_BODY

    # Location line
    p_loc = doc.add_paragraph()
    p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_loc.space_before = Pt(6)
    r_loc = p_loc.add_run(f"Location: {location_name}")
    r_loc.font.size = Pt(12)
    r_loc.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    r_loc.font.name = FONT_BODY

    # Prepared for
    p_cust = doc.add_paragraph()
    p_cust.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cust.space_before = Pt(4)
    r_cust = p_cust.add_run(f"Prepared for: {customer_name}")
    r_cust.font.size = Pt(12)
    r_cust.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    r_cust.font.name = FONT_BODY

    # Prepared by
    p_rep = doc.add_paragraph()
    p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rep.space_before = Pt(18)
    r_rep = p_rep.add_run(f"Prepared by: {rep_name}, {rep_title}")
    r_rep.font.size = Pt(11)
    r_rep.font.name = FONT_BODY

    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.space_before = Pt(4)
    r_date = p_date.add_run(f"Date: {report_date}")
    r_date.font.size = Pt(11)
    r_date.font.name = FONT_BODY

    _add_horizontal_rule(doc)

    # Page break after cover
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # BODY SECTIONS (parsed from Claude output)
    # ═══════════════════════════════════════════════════════════════════════════

    sections = _parse_report_sections(report_content)

    # Section: A Quick Note Before We Start
    intro = sections.get("intro", "")
    if intro:
        _add_section_header(doc, "A Quick Note Before We Start")
        _add_body_paragraph(doc, intro)
        _add_horizontal_rule(doc)

    # Section: What We Found
    findings = sections.get("findings", "")
    if findings:
        _add_section_header(doc, "What We Found")
        # Split by numbered items or paragraphs
        for para in findings.split("\n"):
            para = para.strip()
            if para:
                _add_body_paragraph(doc, para)
        _add_horizontal_rule(doc)

    # Section: Our Recommendations
    recommendations = sections.get("recommendations", "")
    if recommendations:
        _add_section_header(doc, "Our Recommendations")
        for para in recommendations.split("\n"):
            para = para.strip()
            if para:
                _add_body_paragraph(doc, para)
        _add_horizontal_rule(doc)

    # Section: Your Program Summary (pricing table)
    if include_pricing and pricing_table:
        _add_section_header(doc, "Your Program Summary")

        # Build table
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Header row
        headers = ["Product", "Container", "Price/Unit", "Extended Price"]
        for i, header_text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header_text
            _set_cell_shading(cell, "0D1F3C")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = WHITE
                    run.font.name = FONT_BODY

        # Data rows
        for idx, row_data in enumerate(pricing_table):
            row = table.add_row()
            row.cells[0].text = str(row_data.get("product_name", ""))
            row.cells[1].text = str(row_data.get("container", ""))
            row.cells[2].text = f"${row_data.get('price_per_unit', 0):.2f}"
            extended = row_data.get("extended")
            row.cells[3].text = f"${extended:.2f}" if extended else "—"

            # Alternating row shading
            if idx % 2 == 0:
                for cell in row.cells:
                    _set_cell_shading(cell, "F3F4F6")

            # Style data cells
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = FONT_BODY

        # Pricing citation row
        citation_row = table.add_row()
        merged_cell = citation_row.cells[0].merge(citation_row.cells[3])
        merged_cell.text = f"Pricing confirmed for {location_name} as of {report_date}"
        for paragraph in merged_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.font.name = FONT_BODY

        _add_horizontal_rule(doc)

    # Section: What Happens Next
    next_steps = sections.get("next_steps", "")
    if next_steps:
        _add_section_header(doc, "What Happens Next")
        for para in next_steps.split("\n"):
            para = para.strip()
            if para:
                _add_body_paragraph(doc, para)
        _add_horizontal_rule(doc)

    # Section: About Bower Ag
    about = sections.get("about", "")
    if about:
        _add_section_header(doc, "About Bower Ag")
        _add_body_paragraph(doc, about)

    # If parsing produced no sections, dump the full body
    if "body" in sections:
        _add_section_header(doc, "Report")
        for para in sections["body"].split("\n"):
            para = para.strip()
            if para:
                _add_body_paragraph(doc, para)

    # ═══════════════════════════════════════════════════════════════════════════
    # Serialize to bytes
    # ═══════════════════════════════════════════════════════════════════════════

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
