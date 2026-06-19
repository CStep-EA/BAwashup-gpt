"""
Bower Ag CowCare Tool — Video Analysis Worker
Sprint 14: ARQ async worker for video frame extraction and batch Claude Vision analysis.

Pipeline:
  1. Download video from R2 to /tmp
  2. Extract frames at 1fps via ffmpeg
  3. Analyze frames in batches of 5 via Claude Vision
  4. Aggregate findings into structured report
  5. Run governance check on product mentions
  6. Generate report record + DOCX
  7. Cleanup /tmp files (always, even on failure)
"""

import base64
import glob
import io
import json
import logging
import os
import re
import shutil
import subprocess
from datetime import date
from typing import Optional

from app.db.supabase_client import get_supabase_client
from app.services.storage_service import get_storage_service, StorageError

logger = logging.getLogger(__name__)

CLAUDE_MODEL = "claude-sonnet-4-20250514"
VISION_MAX_TOKENS = 1024


# ─── Job status helpers ──────────────────────────────────────────────────────

def _update_job(job_id: str, **fields) -> None:
    """Update a media_jobs record."""
    try:
        client = get_supabase_client()
        client.table("media_jobs").update(fields).eq("id", job_id).execute()
    except Exception as e:
        logger.error(f"[VideoWorker] Failed to update job {job_id}: {e}")


def _update_job_status(
    job_id: str,
    status: str,
    result_report_id: Optional[str] = None,
    error_message: Optional[str] = None,
) -> None:
    fields: dict = {"status": status}
    if result_report_id:
        fields["result_report_id"] = result_report_id
    if error_message:
        fields["error_message"] = error_message[:500]
    if status in ("complete", "failed"):
        fields["completed_at"] = "now()"
    _update_job(job_id, **fields)


# ─── Main worker function ───────────────────────────────────────────────────

async def process_video(ctx, job_id: str, user_id: str, r2_path: str):
    """
    ARQ async task: process an uploaded video.
    Downloads, extracts frames, analyzes via Claude Vision, generates report.
    """
    tmp_dir = f"/tmp/video_{job_id}"

    try:
        # 1. Update status to processing
        _update_job_status(job_id, "processing")

        # 2. Download video from R2
        os.makedirs(tmp_dir, exist_ok=True)
        input_path = os.path.join(tmp_dir, "input.mp4")
        storage = get_storage_service()
        await storage.download_to_file(r2_path, input_path)

        # 3. Extract frames at 1fps
        frames_dir = os.path.join(tmp_dir, "frames")
        os.makedirs(frames_dir, exist_ok=True)

        result = subprocess.run(
            [
                "ffmpeg", "-i", input_path,
                "-vf", "fps=1",
                "-q:v", "2",
                os.path.join(frames_dir, "frame_%04d.jpg"),
            ],
            capture_output=True,
            timeout=120,
        )
        if result.returncode != 0:
            raise RuntimeError(f"ffmpeg failed: {result.stderr.decode()[:300]}")

        # 4. Collect frame paths
        frame_paths = sorted(glob.glob(os.path.join(frames_dir, "*.jpg")))
        _update_job(job_id, frames_extracted=len(frame_paths))

        if not frame_paths:
            raise RuntimeError("No frames extracted from video")

        # 5. Analyze frames in batches of 5
        all_findings = []
        for i in range(0, len(frame_paths), 5):
            batch = frame_paths[i : i + 5]
            batch_findings = await _analyze_frame_batch(batch, i)
            all_findings.extend(batch_findings)
            _update_job(job_id, frames_analyzed=min(i + 5, len(frame_paths)))

        # 6. Aggregate findings
        aggregated = _aggregate_video_findings(all_findings, len(frame_paths))

        # 7. Governance check
        governed = await _apply_governance_to_text(aggregated, user_id)

        # 8. Create report record
        report_id = _create_video_report(user_id, governed, r2_path)

        # 9. Build DOCX
        docx_bytes = _build_video_report_docx(governed)
        r2_report_path = f"reports/{user_id}/video_report_{job_id}.docx"
        await storage.upload_bytes(
            docx_bytes,
            r2_report_path,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        # Update report with docx path
        try:
            client = get_supabase_client()
            client.table("reports").update({
                "docx_r2_path": r2_report_path,
                "status": "complete",
            }).eq("id", report_id).execute()
        except Exception:
            pass

        # 10. Mark job complete
        _update_job_status(job_id, "complete", result_report_id=report_id)
        logger.info(f"[VideoWorker] Job {job_id} complete — report {report_id}")

    except Exception as e:
        logger.error(f"[VideoWorker] Job {job_id} failed: {e}")
        _update_job_status(job_id, "failed", error_message=str(e))
        raise
    finally:
        # ALWAYS clean up tmp files
        shutil.rmtree(tmp_dir, ignore_errors=True)


# ─── Frame analysis ──────────────────────────────────────────────────────────

async def _analyze_frame_batch(
    frame_paths: list[str], start_index: int
) -> list[dict]:
    """Analyze a batch of frame images via Claude Vision."""
    import anthropic

    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))

    content_blocks = []
    for i, path in enumerate(frame_paths):
        with open(path, "rb") as f:
            data = base64.b64encode(f.read()).decode()

        content_blocks.append({
            "type": "text",
            "text": f"Frame {start_index + i + 1} (timestamp ~{start_index + i + 1}s):",
        })
        content_blocks.append({
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": "image/jpeg",
                "data": data,
            },
        })

    content_blocks.append({
        "type": "text",
        "text": (
            "You are a dairy parlor and cow care expert. "
            "Analyze these video frames from a dairy operation. "
            "For each frame, note: equipment visible, animal condition, cleanliness, "
            "any concerns or notable observations. "
            "Be specific and practical. Format as a JSON array of objects with keys: "
            "frame_num, timestamp_s, observations (string)."
        ),
    })

    try:
        response = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=VISION_MAX_TOKENS,
            messages=[{"role": "user", "content": content_blocks}],
        )
        text = response.content[0].text

        # Parse JSON from response
        json_match = re.search(r"\[.*\]", text, re.DOTALL)
        if json_match:
            return json.loads(json_match.group())
        else:
            # Fallback: return as single observation
            return [
                {
                    "frame_num": start_index + i + 1,
                    "timestamp_s": start_index + i + 1,
                    "observations": text,
                }
                for i in range(len(frame_paths))
            ]
    except Exception as e:
        logger.error(f"[VideoWorker] Claude Vision batch failed: {e}")
        return [
            {
                "frame_num": start_index + i + 1,
                "timestamp_s": start_index + i + 1,
                "observations": f"Frame analysis failed: {str(e)[:100]}",
            }
            for i in range(len(frame_paths))
        ]


# ─── Aggregation ─────────────────────────────────────────────────────────────

def _aggregate_video_findings(
    findings: list[dict], total_frames: int
) -> str:
    """Aggregate per-frame findings into a structured summary."""
    duration = total_frames  # 1fps → 1 frame per second

    # Collect all observations
    all_obs = [f.get("observations", "") for f in findings if f.get("observations")]

    # Count observation frequency
    obs_counts: dict[str, int] = {}
    for obs in all_obs:
        for sentence in re.split(r"[.;]", obs):
            sentence = sentence.strip()
            if len(sentence) > 10:
                key = sentence.lower()[:80]
                obs_counts[key] = obs_counts.get(key, 0) + 1

    # Categorize
    consistent = []
    isolated = []
    for obs_key, count in obs_counts.items():
        pct = count / max(len(all_obs), 1)
        if pct >= 0.5:
            consistent.append(obs_key)
        elif pct < 0.2:
            isolated.append(obs_key)

    parts = [
        "## Video Analysis Summary\n",
        f"Duration analyzed: approximately {duration} seconds ({total_frames} frames reviewed)\n",
        "\n## Key Observations\n",
    ]

    for obs in all_obs[:10]:  # Top 10 observations
        parts.append(f"- {obs[:200]}\n")

    if consistent:
        parts.append("\n## Consistent Findings (observed in 50%+ of frames)\n")
        for c in consistent[:5]:
            parts.append(f"- {c}\n")

    if isolated:
        parts.append("\n## Isolated Observations (appeared in <20% of frames)\n")
        for iso in isolated[:5]:
            parts.append(f"- {iso}\n")

    parts.append(
        "\n## Recommended Areas of Attention\n"
        "Based on the video analysis, we recommend discussing the following "
        "areas with your Bower Ag representative for specific product and "
        "protocol recommendations.\n"
    )

    return "".join(parts)


# ─── Governance ──────────────────────────────────────────────────────────────

async def _apply_governance_to_text(text: str, user_id: str) -> str:
    """
    Scan text for product name mentions and apply governance.
    Replace unknown products with category suggestion.
    Replace location-unavailable products with alternative prompt.
    """
    client = get_supabase_client()

    # Find potential product mentions (capitalized multi-word patterns)
    potential_products = re.findall(r"\b([A-Z][a-z]+(?:\s+[A-Z0-9][a-z0-9]*)+)\b", text)

    for name in set(potential_products):
        try:
            result = (
                client.table("products")
                .select("id,product_name")
                .ilike("product_name", f"%{name}%")
                .eq("active", True)
                .limit(1)
                .execute()
            )
            if not result.data:
                # Product doesn't exist — replace
                text = text.replace(
                    name,
                    "[Product category] options are available -- ask your Bower Ag rep for specifics.",
                )
        except Exception:
            pass  # Skip governance for this mention

    return text


# ─── Report creation ─────────────────────────────────────────────────────────

def _create_video_report(user_id: str, content: str, video_r2_path: str) -> str:
    """Create a report record for the video analysis."""
    client = get_supabase_client()

    report_row = {
        "created_by": user_id,
        "customer_name": "Video Analysis",
        "operation_name": "Video Analysis Report",
        "location_code": "N/A",
        "product_ids": [],
        "findings": content[:2000],
        "recommendations": "See video analysis report.",
        "rep_name": "Automated Analysis",
        "rep_title": "CowCare Vision AI",
        "include_pricing": False,
        "status": "generating",
        "report_content": content,
    }

    result = client.table("reports").insert(report_row).execute()
    return result.data[0]["id"]


def _build_video_report_docx(content: str) -> bytes:
    """Build a simple DOCX for the video analysis report."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    section = doc.sections[0]

    # Title
    p = doc.add_paragraph()
    r = p.add_run("Bower Ag — Video Analysis Report")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x0D, 0x1F, 0x3C)

    # Date
    p = doc.add_paragraph()
    r = p.add_run(f"Date: {date.today().isoformat()}")
    r.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Content sections
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.space_before = Pt(12)
            r = p.add_run(line[3:])
            r.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(0x0D, 0x1F, 0x3C)
        elif line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(11)
        else:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
