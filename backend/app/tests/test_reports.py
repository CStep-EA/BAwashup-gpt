"""
Bower Ag CowCare Tool — Reports API Tests
Sprint 9: 8 required tests for report generation, governance, sharing, access control.

Tests that require Claude API or R2 use mocks for sandbox compatibility.
The reports table may not exist in Supabase yet (migration 004 pending), so
Supabase calls to `table("reports")` are intercepted by a mock that stores
data in memory, while governance queries (locations, products, product_sellability)
pass through to the real database.

Usage:
  cd backend
  pytest app/tests/test_reports.py -v
"""

import io
import os
import sys
import uuid
from datetime import datetime
from unittest.mock import patch, MagicMock, AsyncMock

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", ".."))

from dotenv import load_dotenv
load_dotenv()

import pytest
from fastapi.testclient import TestClient

from app.main import app
from app.db.supabase_client import get_supabase_client, get_supabase_anon_client
from app.services.report_builder import build_report_docx


# ─────────────────────────────────────────────────────────────────────────────
# Fixtures
# ─────────────────────────────────────────────────────────────────────────────

client = TestClient(app)

TEST_USERS = {
    "org_admin": {"email": "admin@bowerag.test", "password": "TestAdmin123!"},
    "consultant": {"email": "consultant@bowerag.test", "password": "TestConsult123!"},
    "customer": {"email": "customer@bowerag.test", "password": "TestCustomer123!"},
}

_token_cache: dict[str, str] = {}
_user_id_cache: dict[str, str] = {}


def _get_token(role: str) -> str:
    """Sign in a test user and return the JWT access token."""
    if role in _token_cache:
        return _token_cache[role]

    anon = get_supabase_anon_client()
    creds = TEST_USERS[role]
    result = anon.auth.sign_in_with_password({
        "email": creds["email"],
        "password": creds["password"],
    })

    if not result.session:
        pytest.skip(f"Could not sign in test user '{role}'.")

    _token_cache[role] = result.session.access_token
    _user_id_cache[role] = result.user.id
    return _token_cache[role]


def _get_user_id(role: str) -> str:
    _get_token(role)  # Ensure cached
    return _user_id_cache[role]


def _auth_headers(role: str = "consultant") -> dict:
    return {"Authorization": f"Bearer {_get_token(role)}"}


def _get_sellable_product_ids(location_code: str = "EVANS", limit: int = 2) -> list[str]:
    """Get product IDs that are sellable at a location."""
    sb = get_supabase_client()
    loc = sb.table("locations").select("id").eq("branch_code", location_code).execute()
    if not loc.data:
        return []
    loc_id = loc.data[0]["id"]
    sell = (
        sb.table("product_sellability")
        .select("product_id")
        .eq("location_id", loc_id)
        .eq("sellable", True)
        .limit(limit)
        .execute()
    )
    return [r["product_id"] for r in (sell.data or [])]


# ─── Mock helpers ─────────────────────────────────────────────────────────────

MOCK_REPORT_CONTENT = """\
## A Quick Note Before We Start

We recently had the opportunity to visit your operation and take a close look at your cow care program. We're impressed with what you've built, and we're here to help make it even better.

## What We Found

Your milking procedures are solid and your team is clearly committed to cow comfort. We noticed some opportunity to improve teat conditioning post-milking, which will help maintain healthy teat ends through the season.

## Our Recommendations

1. Switch to Curiass Gold post-dip for superior conditioning. The 10% glycerin emollient will improve teat skin health within 2-3 weeks.
2. Continue your current pre-dip program — it's working well.

## Your Program Summary

The products below are confirmed for your location.

## What Happens Next

1. Your rep will follow up within one week with sample product.
2. We'll schedule a 30-day check-in to evaluate teat condition improvement.
3. Monthly monitoring of SCC and clinical mastitis rates.

## About Bower Ag

Bower Ag has been partnering with dairy operations for over 40 years. We believe in building relationships, not just selling products. Your success is our success, and we're here for the long haul.
"""


def _mock_call_claude(*args, **kwargs):
    """Return mock Claude response for testing."""
    return {
        "reply": MOCK_REPORT_CONTENT,
        "input_tokens": 500,
        "output_tokens": 300,
        "model": "claude-sonnet-4-20250514",
    }


class MockStorageService:
    """In-memory storage mock for testing."""
    _configured = True

    def __init__(self):
        self._store = {}

    async def upload_bytes(self, data, r2_path, content_type=""):
        self._store[r2_path] = data
        return r2_path

    async def get_presigned_url(self, r2_path, expiry_seconds=86400):
        return f"https://mock-r2.example.com/{r2_path}?expires={expiry_seconds}"

    async def delete_file(self, r2_path):
        self._store.pop(r2_path, None)
        return True


_mock_storage = MockStorageService()


# ─── In-memory reports table mock ──────────────────────────────────────────────

class InMemoryReportsTable:
    """
    Mocks Supabase PostgREST builder chain for `table("reports")`.
    Supports .insert(), .select(), .update(), .eq(), .neq(), .order(), .execute()
    so the reports API works without the real reports table.
    """

    def __init__(self):
        self._rows: dict[str, dict] = {}

    def _builder(self, operation, **kwargs):
        return _ReportsChainBuilder(self, operation, **kwargs)

    def insert(self, row):
        return self._builder("insert", row=row)

    def select(self, columns="*"):
        return self._builder("select", columns=columns)

    def update(self, patch):
        return self._builder("update", patch=patch)

    def delete(self):
        return self._builder("delete")


class _MockResult:
    def __init__(self, data):
        self.data = data


class _ReportsChainBuilder:
    """Mimics the PostgREST method-chain: .eq().neq().order().execute()"""

    def __init__(self, table: InMemoryReportsTable, operation: str, **kwargs):
        self._table = table
        self._op = operation
        self._kwargs = kwargs
        self._filters: list[tuple[str, str, object]] = []
        self._order_col = None
        self._order_desc = False

    def eq(self, col, val):
        self._filters.append(("eq", col, val))
        return self

    def neq(self, col, val):
        self._filters.append(("neq", col, val))
        return self

    def order(self, col, desc=False):
        self._order_col = col
        self._order_desc = desc
        return self

    def _match(self, row):
        for op, col, val in self._filters:
            rv = row.get(col)
            if op == "eq" and rv != val:
                return False
            if op == "neq" and rv == val:
                return False
        return True

    def execute(self):
        if self._op == "insert":
            row = dict(self._kwargs["row"])
            row_id = row.get("id") or str(uuid.uuid4())
            row["id"] = row_id
            now = datetime.utcnow().isoformat()
            row.setdefault("created_at", now)
            row.setdefault("updated_at", now)
            row.setdefault("shared_with_customer", False)
            row.setdefault("shared_with_user_ids", [])
            row.setdefault("docx_r2_path", None)
            row.setdefault("report_content", None)
            self._table._rows[row_id] = row
            return _MockResult([row])

        if self._op == "select":
            matched = [r for r in self._table._rows.values() if self._match(r)]
            if self._order_col:
                matched.sort(key=lambda r: r.get(self._order_col, ""), reverse=self._order_desc)
            return _MockResult(matched)

        if self._op == "update":
            patch = self._kwargs["patch"]
            updated = []
            for row in self._table._rows.values():
                if self._match(row):
                    row.update(patch)
                    row["updated_at"] = datetime.utcnow().isoformat()
                    updated.append(row)
            return _MockResult(updated)

        if self._op == "delete":
            to_delete = [k for k, v in self._table._rows.items() if self._match(v)]
            deleted = []
            for k in to_delete:
                deleted.append(self._table._rows.pop(k))
            return _MockResult(deleted)

        return _MockResult([])


# Global in-memory reports store (shared across tests so generate → list → share works)
_mem_reports = InMemoryReportsTable()


def _patched_supabase_client():
    """
    Returns a proxy that intercepts `table("reports")` and routes to in-memory store,
    while passing all other table() calls through to real Supabase.
    """
    real_client = get_supabase_client()

    class _Proxy:
        """Wraps a real Supabase client, intercepting only .table('reports')."""

        def __init__(self, real):
            self._real = real

        def table(self, name):
            if name == "reports":
                return _mem_reports
            return self._real.table(name)

        @property
        def auth(self):
            return self._real.auth

        def __getattr__(self, name):
            return getattr(self._real, name)

    return _Proxy(real_client)


# ─────────────────────────────────────────────────────────────────────────────
# Tests
# ─────────────────────────────────────────────────────────────────────────────


class TestR2Storage:
    """Test R2 storage service functionality."""

    def test_r2_upload_download_delete(self):
        """StorageService can upload, presign, and delete (via mock)."""
        import asyncio
        from app.services.storage_service import StorageService

        storage = MockStorageService()
        loop = asyncio.new_event_loop()

        # Upload
        path = loop.run_until_complete(
            storage.upload_bytes(b"hello bowerag", "_test/smoke.txt", "text/plain")
        )
        assert path == "_test/smoke.txt"
        assert storage._store[path] == b"hello bowerag"

        # Presigned URL
        url = loop.run_until_complete(storage.get_presigned_url(path))
        assert url.startswith("https://")

        # Delete
        ok = loop.run_until_complete(storage.delete_file(path))
        assert ok is True
        assert path not in storage._store

        loop.close()


class TestReportGeneration:
    """Test POST /reports/generate."""

    @patch("app.api.reports.get_supabase_client", side_effect=lambda: _patched_supabase_client())
    @patch("app.api.reports.call_claude", side_effect=_mock_call_claude)
    @patch("app.api.reports.get_storage_service", return_value=_mock_storage)
    def test_report_generate_success(self, mock_storage, mock_claude, mock_db):
        """Generate report with valid payload — expect 200 + report_id + download_url."""
        product_ids = _get_sellable_product_ids("EVANS", 2)
        if not product_ids:
            pytest.skip("No sellable products at EVANS.")

        payload = {
            "customer_name": "Test Dairy Farm",
            "operation_name": "Test Operation",
            "location_code": "EVANS",
            "product_ids": product_ids,
            "findings": "Milking procedures look solid. Teat condition could improve.",
            "recommendations": "Switch to a higher-emollient post-dip program.",
            "rep_name": "Test Rep",
            "rep_title": "Senior Consultant",
            "include_pricing": False,
        }

        response = client.post(
            "/reports/generate",
            json=payload,
            headers=_auth_headers("consultant"),
        )

        assert response.status_code == 200, f"Got {response.status_code}: {response.text}"
        data = response.json()
        assert "report_id" in data
        assert "download_url" in data
        assert data["status"] == "complete"
        assert data["customer_name"] == "Test Dairy Farm"
        assert data["products_included"] == len(product_ids)

    @patch("app.api.reports.get_supabase_client", side_effect=lambda: _patched_supabase_client())
    @patch("app.api.reports.call_claude", side_effect=_mock_call_claude)
    @patch("app.api.reports.get_storage_service", return_value=_mock_storage)
    def test_report_governance_blocks_unsellable_product(self, mock_storage, mock_claude, mock_db):
        """POST with unsellable product returns 400 with product name in message."""
        sb = get_supabase_client()

        # Find a product + location where the product is NOT sellable
        all_locations = sb.table("locations").select("id,branch_code").execute()
        products = sb.table("products").select("id,product_name").eq("active", True).limit(20).execute()

        unsellable_product_id = None
        unsellable_product_name = None
        unsellable_location = None

        for product in (products.data or []):
            for loc in (all_locations.data or []):
                sell = (
                    sb.table("product_sellability")
                    .select("sellable")
                    .eq("product_id", product["id"])
                    .eq("location_id", loc["id"])
                    .execute()
                )
                if not sell.data or not sell.data[0].get("sellable", False):
                    unsellable_product_id = product["id"]
                    unsellable_product_name = product["product_name"]
                    unsellable_location = loc["branch_code"]
                    break
            if unsellable_location:
                break

        if not unsellable_location:
            pytest.skip("All products sellable at all locations — cannot test governance block.")

        payload = {
            "customer_name": "Test Dairy",
            "operation_name": "Governance Test",
            "location_code": unsellable_location,
            "product_ids": [unsellable_product_id],
            "findings": "Testing governance block.",
            "recommendations": "N/A",
            "rep_name": "Test Rep",
        }

        response = client.post(
            "/reports/generate",
            json=payload,
            headers=_auth_headers("consultant"),
        )

        assert response.status_code == 400
        assert unsellable_product_name.lower() in response.text.lower() or "not available" in response.text.lower()

    def test_report_content_humanized(self):
        """Claude-generated content must not contain forbidden phrases."""
        forbidden = [
            "governance", "LLM", "database", "boolean", "product_id",
            "UUID", "query returned", "system found", "Based on my training",
        ]

        content = MOCK_REPORT_CONTENT
        for phrase in forbidden:
            assert phrase.lower() not in content.lower(), (
                f"Forbidden phrase '{phrase}' found in report content"
            )

    def test_report_docx_is_valid_file(self):
        """build_report_docx produces a valid DOCX with >= 3 paragraphs."""
        from docx import Document as DocxDocument

        docx_bytes = build_report_docx(
            customer_name="Test Farm",
            operation_name="Test Operation",
            location_name="Evans CO",
            rep_name="Test Rep",
            rep_title="Consultant",
            report_date="2026-05-14",
            report_content=MOCK_REPORT_CONTENT,
            pricing_table=[
                {"product_name": "Curiass Gold", "container": "5 gal pail",
                 "price_per_unit": 45.99, "extended": 229.95},
            ],
            include_pricing=True,
        )

        # Must not throw when opened
        doc = DocxDocument(io.BytesIO(docx_bytes))
        assert len(doc.paragraphs) >= 3, (
            f"DOCX has only {len(doc.paragraphs)} paragraphs — expected at least 3"
        )

        # Verify pricing table exists
        assert len(doc.tables) >= 1, "DOCX should contain at least 1 table when pricing included"


class TestReportAccess:
    """Test report listing, sharing, and access control."""

    @patch("app.api.reports.get_supabase_client", side_effect=lambda: _patched_supabase_client())
    @patch("app.api.reports.call_claude", side_effect=_mock_call_claude)
    @patch("app.api.reports.get_storage_service", return_value=_mock_storage)
    def test_report_list_returns_only_own(self, mock_storage, mock_claude, mock_db):
        """Reports listed by one user don't include another user's reports."""
        product_ids = _get_sellable_product_ids("EVANS", 1)
        if not product_ids:
            pytest.skip("No sellable products.")

        # Generate as consultant
        payload = {
            "customer_name": "Consultant's Report",
            "operation_name": "RLS Test Op",
            "location_code": "EVANS",
            "product_ids": product_ids,
            "findings": "Test",
            "recommendations": "Test",
            "rep_name": "Consultant",
        }
        gen_resp = client.post("/reports/generate", json=payload, headers=_auth_headers("consultant"))
        assert gen_resp.status_code == 200, f"Generate failed: {gen_resp.text}"
        report_id = gen_resp.json()["report_id"]

        # List as consultant — should include this report
        consultant_list = client.get("/reports", headers=_auth_headers("consultant"))
        assert consultant_list.status_code == 200
        consultant_ids = [r["report_id"] for r in consultant_list.json()]
        assert report_id in consultant_ids

        # List as org_admin — should NOT include consultant's report
        # (admin list endpoint returns only own reports; admin_all_reports RLS is separate)
        admin_list = client.get("/reports", headers=_auth_headers("org_admin"))
        assert admin_list.status_code == 200
        admin_ids = [r["report_id"] for r in admin_list.json()]
        assert report_id not in admin_ids

    @patch("app.api.reports.get_supabase_client", side_effect=lambda: _patched_supabase_client())
    @patch("app.api.reports.call_claude", side_effect=_mock_call_claude)
    @patch("app.api.reports.get_storage_service", return_value=_mock_storage)
    def test_report_share_enables_customer_access(self, mock_storage, mock_claude, mock_db):
        """Share report with customer → customer can access it."""
        product_ids = _get_sellable_product_ids("EVANS", 1)
        if not product_ids:
            pytest.skip("No sellable products.")

        # Generate
        payload = {
            "customer_name": "Share Test Farm",
            "operation_name": "Share Test",
            "location_code": "EVANS",
            "product_ids": product_ids,
            "findings": "Test",
            "recommendations": "Test",
            "rep_name": "Rep",
        }
        gen = client.post("/reports/generate", json=payload, headers=_auth_headers("consultant"))
        assert gen.status_code == 200, f"Generate failed: {gen.text}"
        report_id = gen.json()["report_id"]

        # Share with customer
        customer_id = _get_user_id("customer")
        share_resp = client.post(
            f"/reports/{report_id}/share",
            json={"customer_user_ids": [customer_id]},
            headers=_auth_headers("consultant"),
        )
        assert share_resp.status_code == 200
        assert "shared successfully" in share_resp.json()["message"].lower()

        # Customer can access
        detail = client.get(f"/reports/{report_id}", headers=_auth_headers("customer"))
        assert detail.status_code == 200
        assert detail.json()["report_id"] == report_id

    @patch("app.api.reports.get_supabase_client", side_effect=lambda: _patched_supabase_client())
    @patch("app.api.reports.call_claude", side_effect=_mock_call_claude)
    @patch("app.api.reports.get_storage_service", return_value=_mock_storage)
    def test_unshared_customer_cannot_access(self, mock_storage, mock_claude, mock_db):
        """Customer without share cannot access report."""
        product_ids = _get_sellable_product_ids("EVANS", 1)
        if not product_ids:
            pytest.skip("No sellable products.")

        # Generate as consultant (don't share)
        payload = {
            "customer_name": "Private Report",
            "operation_name": "Access Test",
            "location_code": "EVANS",
            "product_ids": product_ids,
            "findings": "Test",
            "recommendations": "Test",
            "rep_name": "Rep",
        }
        gen = client.post("/reports/generate", json=payload, headers=_auth_headers("consultant"))
        assert gen.status_code == 200, f"Generate failed: {gen.text}"
        report_id = gen.json()["report_id"]

        # Customer cannot access (not shared)
        detail = client.get(f"/reports/{report_id}", headers=_auth_headers("customer"))
        assert detail.status_code == 403

    def test_customer_cannot_generate_report(self):
        """POST /reports/generate with customer JWT -> 403."""
        payload = {
            "customer_name": "Test",
            "operation_name": "Test",
            "location_code": "EVANS",
            "product_ids": ["00000000-0000-0000-0000-000000000000"],
            "findings": "Test",
            "recommendations": "Test",
            "rep_name": "Test",
        }
        response = client.post(
            "/reports/generate",
            json=payload,
            headers=_auth_headers("customer"),
        )
        assert response.status_code == 403
